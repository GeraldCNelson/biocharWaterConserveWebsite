from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Inches

SCRIPT_PATH = Path(__file__).resolve()
BIOCHAR_APP_DIR = SCRIPT_PATH.parents[2]

BASE = BIOCHAR_APP_DIR / "data-processed" / "management" / "irrigation" / "analysis"
DIAGNOSTICS_DIR = BASE / "diagnostics"
FIGURES_DIR = BASE / "figures"
REPORTS_DIR = BASE / "reports"


def _fmt_minutes(value: object) -> str:
    if pd.isna(value):
        return "NA"
    return f"{float(value):.0f}"


def _fmt_datetime(value: object) -> str:
    if pd.isna(value):
        return "NA"

    ts = pd.to_datetime(value, errors="coerce")
    if pd.isna(ts):
        return "NA"

    return pd.Timestamp(ts).strftime("%Y-%m-%d %H:%M")


def _fmt_bool_flag(value: object) -> bool:
    if isinstance(value, bool):
        return value

    if pd.isna(value):
        return False

    return str(value).strip().lower() in {"true", "1", "yes", "y"}


def add_arrival_definitions_section(doc: Document) -> None:
    doc.add_heading("Arrival Detection Definitions", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Definition"
    hdr[2].text = "Purpose"

    row = table.add_row().cells
    row[0].text = "Standard arrival"
    row[1].text = (
        "First time after recorded irrigation start that VWC exceeds baseline "
        "by +0.25% and remains elevated for at least 1 hour "
        "(4 consecutive 15-minute points)."
    )
    row[2].text = (
        "Primary arrival estimate tied to the recorded irrigation start time."
    )

    row = table.add_row().cells
    row[0].text = "Alternate arrival"
    row[1].text = (
        "First VWC step increase of at least +0.50% that is followed by a "
        "sustained rise in VWC. This search is allowed to detect responses "
        "before the recorded irrigation start."
    )
    row[2].text = (
        "QA check for possible early wetting, timing issues, or stronger wetting fronts."
    )

    row = table.add_row().cells
    row[0].text = "Plateau VWC"
    row[1].text = (
        "Post-irrigation stabilized VWC value estimated from the event response."
    )
    row[2].text = (
        "Used later for estimating water stored in the sensor zone."
    )

def build_arrival_diagnostics_report(
    year: int,
    include_expected: bool = False,
    output_path: Optional[str | Path] = None,
) -> Path:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    arrival_order = pd.read_csv(
        DIAGNOSTICS_DIR / f"arrival_order_diagnostics_{year}.csv"
    )

    arrival_times = pd.read_csv(
        DIAGNOSTICS_DIR / f"irrigation_arrival_times_{year}.csv"
    )
    plot_log = pd.read_csv(
        FIGURES_DIR / f"irrigation_event_multidepth_plot_log_{year}.csv"
    )

    # One row per event/strip/logger with event timing and thresholds.
    event_meta = (
        arrival_times[
            [
                "event_id",
                "strip",
                "logger_position",
                "irrigation_start",
                "irrigation_end",
                "event_duration_hours",
                "gallons_strip",
                "arrival_threshold_vwc",
                "alt_arrival_threshold_vwc",
            ]
        ]
        .drop_duplicates()
        .copy()
    )

    report_rows = arrival_order.copy()

    if not include_expected:
        report_rows = report_rows[
            report_rows["order_class"].ne("as expected")
            | report_rows["alt_order_class"].ne("as expected")
            | report_rows["any_alt_before_start"].apply(_fmt_bool_flag)
        ].copy()

    report_rows = report_rows.merge(
        event_meta,
        on=["event_id", "strip", "logger_position"],
        how="left",
    )

    report_rows = report_rows.merge(
        plot_log[["event_id", "strip", "logger_position", "output_file", "status"]],
        on=["event_id", "strip", "logger_position"],
        how="left",
    )

    report_rows = report_rows[report_rows["status"].eq("written")].copy()

    report_rows["irrigation_start_sort"] = pd.to_datetime(
        report_rows["irrigation_start"],
        errors="coerce",
    )

    report_rows["event_date_from_id"] = pd.to_datetime(
        report_rows["event_id"].astype(str).str.slice(0, 10),
        errors="coerce",
    )

    report_rows["irrigation_start_sort"] = pd.to_datetime(
        report_rows["irrigation_start"],
        errors="coerce",
    )

    report_rows["calendar_sort"] = report_rows["irrigation_start_sort"].fillna(
        report_rows["event_date_from_id"]
    )

    report_rows = report_rows.sort_values(
        ["calendar_sort", "event_id", "strip", "logger_position"],
        kind="stable",
    ).reset_index(drop=True)

    doc = Document()
    doc.add_heading(f"{year} Irrigation Arrival Diagnostics", level=1)
    doc.add_paragraph(
        "Comparison of standard sustained-response arrival detection and "
        "alternate sustained step-change arrival detection. Plateau VWC is retained "
        "because it is used for estimating stored water after irrigation."
    )

    add_arrival_definitions_section(doc)

    if include_expected:
        doc.add_paragraph("Includes all arrival-order classes.")
    else:
        doc.add_paragraph(
            "Includes only events with unexpected primary arrival order, "
            "unexpected alternate arrival order, or alternate arrivals before "
            "the recorded irrigation start."
        )

    if report_rows.empty:
        doc.add_paragraph("No matching plots found.")
    else:
        for _, row in report_rows.iterrows():
            doc.add_heading(
                (
                    f"{_fmt_datetime(row.get('irrigation_start'))} | "
                    f"{row['strip']} {row['logger_position']} | "
                    f"primary: {row['order_class']} | "
                    f"alt: {row['alt_order_class']}"
                ),
                level=2,
            )

            doc.add_paragraph(f"Event ID: {row['event_id']}")
            doc.add_paragraph(
                f"Irrigation start: {_fmt_datetime(row.get('irrigation_start'))} | "
                f"Irrigation end: {_fmt_datetime(row.get('irrigation_end'))} | "
                f"Duration: {_fmt_minutes(row.get('event_duration_hours'))} hr | "
                f"Strip volume: {_fmt_minutes(row.get('gallons_strip'))} gal"
            )

            # doc.add_paragraph(
            #     "Standard arrival "
            #     f"(+{row.get('arrival_threshold_vwc', 0.25)}% VWC over baseline; "
            #     "first sustained response after recorded irrigation start):"
            # )

            doc.add_paragraph(
                f"Standard arrival, order: {row['arrival_order']} | "
                f"6 in: {_fmt_minutes(row['arrival_6in_min'])} min | "
                f"12 in: {_fmt_minutes(row['arrival_12in_min'])} min | "
                f"18 in: {_fmt_minutes(row['arrival_18in_min'])} min"
            )

            # doc.add_paragraph(
            #     "Alternate arrival "
            #     f"(+{row.get('alt_arrival_threshold_vwc', 0.5)}% VWC step; "
            #     "first sustained VWC rise, allowed before recorded irrigation start):"
            # )

            doc.add_paragraph(
                f"Alternate arrival, order: {row['alt_arrival_order']} | "
                f"6 in: {_fmt_minutes(row['alt_arrival_6in_min'])} min | "
                f"12 in: {_fmt_minutes(row['alt_arrival_12in_min'])} min | "
                f"18 in: {_fmt_minutes(row['alt_arrival_18in_min'])} min"
            )

            alt_before_depths = row.get("alt_before_start_depths", "")
            if pd.notna(alt_before_depths) and str(alt_before_depths).strip():
                doc.add_paragraph(
                    "⚠ Alternate arrival detected before recorded irrigation start "
                    f"at depths: {alt_before_depths} in"
                )

            image_path = Path(str(row["output_file"]))
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.5))
            else:
                doc.add_paragraph(f"Missing figure: {image_path}")

    if output_path is None:
        suffix = "all" if include_expected else "unexpected"
        output_path = REPORTS_DIR / f"irrigation_arrival_{suffix}_{year}.docx"
    else:
        output_path = Path(output_path)

    doc.save(output_path)
    return output_path


def build_holding_capacity_report() -> None:
    raise NotImplementedError("Holding capacity report is not implemented yet.")


def build_trustworthy_events_report() -> None:
    raise NotImplementedError("Trustworthy events report is not implemented yet.")


def main() -> None:
    year = 2026

    output_path = build_arrival_diagnostics_report(
        year=year,
        include_expected=False,
    )

    print(f"Wrote report: {output_path}")


if __name__ == "__main__":
    main()